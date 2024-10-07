# app.py

import streamlit as st
import boto3
import json
import os
from io import BytesIO
# Library for PDF text extraction
import PyPDF2
# Library for creating Word documents
from docx import Document
from docx.shared import Pt


# Set page configuration with custom title and favicon
st.set_page_config(
    page_title="Resume Builder",
    page_icon="NIWC_Atlantic_Logo.jpg"
)

# Initialize AWS clients
bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')
s3 = boto3.client('s3')

def get_rag_data_from_pdf():
    bucket_name = 'resume-builder-mockup-bucket'
    keys = ['Federal Resume Samples.pdf', 'sample-resume.pdf']
    combined_text = ''

    for key in keys:
        try:
            response = s3.get_object(Bucket=bucket_name, Key=key)
            pdf_content = response['Body'].read()
            pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_content))
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text()
            combined_text += text + '\n'  # Append text from each PDF with a separator
        except Exception as e:
            st.error(f"Error fetching or parsing PDF from S3 for {key}: {e}")
            continue  # Skip to the next key in case of an error
    
    return combined_text

def create_prompt(experience):
    rag_data = get_rag_data_from_pdf()
    prompt = f"""
Human:
Here's a good bit of information about resumes generally speaking and a few good examples of how resumes should read {rag_data}


Based on the following information provided under 'Work Experience:', generate the 'Experience' section of a professional resume. Reword and enhance upon the details to make them more compelling and suitable for inclusion in a professional resume.


Please format the output as follows:


### Experience ###
(Enhanced experience content here)


Do not include any personal information like name, contact details, or education.
Do not include any new numbers and percentages in terms of impact. The only quantification of impact should be whatever the user literally states.
If the user only gives limited input in terms of experience, only reword and slightly enhance that experience rather than adding a great amount of detail.

Work Experience:
{experience}
Assistant:
"""

    return prompt

def generate_experience(prompt):
    try:
        payload = {
            "prompt": prompt
        }
        request_body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 50000,
            "temperature": 0.7,
            "messages": [
                {"role": "user", "content": prompt}
            ]
        })
        response = bedrock.invoke_model(
            modelId='anthropic.claude-3-5-sonnet-20240620-v1:0',
            contentType='application/json',
            accept='application/json',
            body=request_body
        )
        response_body = response['body'].read()
        result = json.loads(response_body.decode('utf-8'))

        # Debugging: Print or display the result
        st.write("Model Response:", result)

        # Extract the generated text from the response
        content = result.get('content', [])
        if content and isinstance(content, list) and 'text' in content[0]:
            generated_text = content[0]['text']
        else:
            st.error("Unexpected response format from the model.")
            return ""

        return generated_text

    except Exception as e:
        st.error(f"An error occurred while invoking the model: {e}")
        return ""

def parse_experience(experience_text):
    # Parse the AI-generated experience section
    lines = experience_text.splitlines()
    experience_content = ''
    start_parsing = False
    for line in lines:
        line = line.strip()
        if line == "### Experience ###":
            start_parsing = True
            continue
        elif start_parsing:
            experience_content += line + '\n'
    return experience_content.strip()

def format_skills(skills_input):
    # Format the skills as per user's request
    return skills_input.strip()

def create_resume_docx(full_name, email, phone, education, experience_content, skills_content):
    # Create a new Word document
    document = Document()

    # Add Name as a header
    name_paragraph = document.add_heading(full_name, level=0)
    name_paragraph.alignment = 1  # Center alignment

    # Add contact information
    contact_info = f"Email: {email}\nPhone: {phone}"
    contact_paragraph = document.add_paragraph(contact_info)
    contact_paragraph.alignment = 1  # Center alignment

    # Add Education section
    document.add_heading('Education', level=1)
    document.add_paragraph(education)

    # Add Experience section
    document.add_heading('Experience', level=1)
    document.add_paragraph(experience_content)

    # Add Skills section
    document.add_heading('Skills', level=1)
    skills_lines = skills_content.split('\n')
    for line in skills_lines:
        # Each line is a separate category
        line = "- " + line
        document.add_paragraph(line)

    # Save the document
    docx_io = BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io

def main():
    st.title("NIWC-A AI Resume Builder Mockup")

    st.header("Personal Information")
    full_name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    education = st.text_area("Education (e.g., Degree, University)")

    st.header("Work Experience")
    experience = st.text_area("Describe your work experience")

    st.header("Skills")
    skills = st.text_area("List your skills as per the format:\n- Please group together similar skills on one line and separate them with a comma.\n\n For separate skill categories, write on a new line.\n\nExample:\n\n\'Python, Java, R\'\n\n\'Leadership, Communication\'")

    submit = st.button("Generate Resume")

    if submit:
        if not all([full_name, email, phone, education, experience, skills]):
            st.error("Please fill in all the fields.")
        else:
            with st.spinner('Generating your resume...'):
                prompt = create_prompt(experience)
                experience_text = generate_experience(prompt)

            if experience_text:
                # Parse the AI-generated experience content
                experience_content = parse_experience(experience_text)

                if not experience_content:
                    st.error("Failed to parse the AI-generated content. Please try again.")
                else:
                    # Format the skills as per user's request
                    skills_content = format_skills(skills)

                    # Create the DOCX resume
                    docx_file = create_resume_docx(full_name, email, phone, education, experience_content, skills_content)

                    st.header("Your Generated Resume")
                    st.download_button(
                        'Download Resume',
                        data=docx_file,
                        file_name='resume.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    st.success("Your resume has been generated!")

                    # Display the text version of the resume
                    text_resume = f"""
Name: {full_name}
Email: {email}
Phone: {phone}
Education: {education}

Experience:
{experience_content}

Skills:
{skills_content}
"""
                    st.text_area("Text Version of Your Resume", text_resume, height=300)

if __name__ == "__main__":
    main()
